import sys
import subprocess

def install_docx():
    try:
        import docx
    except ImportError:
        print("Installing python-docx...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        print("Installed python-docx successfully.")

install_docx()

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_quotation():
    doc = Document()
    
    # Title
    title = doc.add_heading('Commercial Quotation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Header Info
    doc.add_paragraph('Project Name: Comprehensive School Management System (SaaS Platform)')
    doc.add_paragraph('Date: April 28, 2026')
    doc.add_paragraph('Prepared For: Navadaya / Nabodaya Educational Trust')
    
    doc.add_heading('1. Project Overview', level=1)
    p = doc.add_paragraph(
        'A comprehensive, multi-tenant School Management System designed to automate administrative, '
        'academic, and financial operations. The system supports multiple organizations and schools '
        'under a single unified dashboard, complete with public-facing websites and secure role-based portals.'
    )
    
    doc.add_heading('2. Key Features & Modules', level=1)
    features = [
        "Multi-Tenant Architecture (Super Admin, Organization Admin, School Admin)",
        "Role-based Access Control (Teachers, Students, Parents, Accountants)",
        "Student & Teacher Management (Profiles, Records, Documents)",
        "Admissions Management (Online applications, Status tracking, Auto roll-number assignment)",
        "Academic Management (Classes, Subjects, Assignments, Timetable)",
        "Attendance System (Daily tracking for students and staff)",
        "Fees & Financial Management (Invoicing, Collections, Receipts)",
        "Examinations Module (Exam scheduling, Grading, Report Cards)",
        "Library Management (Book inventory, Issuing, Returns)",
        "Inventory Management (Stock tracking, Issue to staff/students)",
        "Live Classes Integration (Virtual learning environment)",
        "Mobile Application (Android & iOS) for Students, Parents & Teachers",
        "Communication & Notifications (Email, SMS & Push alerts)",
        "Public Landing Pages (Customizable school websites)",
        "Reporting & Analytics (Graphical dashboards and performance metrics)"
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
        
    doc.add_heading('3. Technology Stack', level=1)
    doc.add_paragraph('Backend: Python, Django REST Framework, PostgreSQL, Redis, Celery', style='List Bullet')
    doc.add_paragraph('Frontend: React, Vite, Tailwind CSS, TypeScript', style='List Bullet')
    doc.add_paragraph('Deployment: Docker, Coolify, Traefik, Ubuntu VPS', style='List Bullet')
    
    doc.add_heading('4. Commercials (Estimated)', level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item / Phase'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Cost (INR)'
    
    costs = [
        ("Phase 1: Foundation & Core Backend", "Multi-tenant architecture, role-based RBAC, core models, and Docker configuration.", "₹ 35,000"),
        ("Phase 2: Core Admin Modules", "Students, Teachers, Classes, Sections, and Subjects CRUD with professional UI.", "₹ 40,000"),
        ("Phase 3: Financial & Academic", "Fee management, Razorpay integration, Auto-grading, and PDF Report Cards.", "₹ 30,000"),
        ("Phase 4: Attendance & Communication", "Leave management workflow, automated notifications (Email/SMS structure).", "₹ 20,000"),
        ("Phase 5: Live Classes & Assignments", "Virtual classrooms (Zoom/Meet), Assignment submissions, and Teacher feedback.", "₹ 25,000"),
        ("Phase 6: Mobile App Development", "Cross-platform React Native app for Students & Parents (Attendance, Fees, Dashboard).", "₹ 50,000"),
        ("Phase 7: Analytics & Dashboards", "Advanced reporting, visual analytics, and public landing pages.", "₹ 15,000"),
        ("Deployment & Infrastructure", "Coolify setup, SSL, Domain routing, and VPS configuration.", "₹ 15,000"),
        ("Total Estimated Project Cost", "Complete source code, mobile app, and production-ready deployment.", "₹ 2,30,000"),
        ("Annual Maintenance (Optional)", "Server monitoring, security patches, and routine updates.", "₹ 35,000 / year")
    ]
    
    for item, desc, cost in costs:
        row_cells = table.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = desc
        row_cells[2].text = cost
        
    doc.add_paragraph('\n* Note: Above costs are placeholder estimates. Final billing depends on mutual agreement and specific feature requirements.')
    
    doc.add_heading('5. Terms and Conditions', level=1)
    doc.add_paragraph('1. Payment Terms: 30% advance, 40% upon beta delivery, 30% upon final deployment.', style='List Number')
    doc.add_paragraph('2. Timeline: Estimated 8-12 weeks for complete delivery.', style='List Number')
    doc.add_paragraph('3. Intellectual Property: Source code ownership will be transferred upon full payment.', style='List Number')
    doc.add_paragraph('4. Hosting: VPS and domain costs are to be borne by the client directly.', style='List Number')
    
    doc.add_paragraph('\n----------------------------------------\nAuthorized Signature')
    
    filename = 'School_Management_System_Quotation_V3.docx'
    doc.save(filename)
    print(f"Successfully created {filename}")

if __name__ == '__main__':
    create_quotation()
